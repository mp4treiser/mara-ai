import os
from abc import ABC, abstractmethod
from typing import List
import logging

logger = logging.getLogger(__name__)


class FileHandler(ABC):
    """Базовый класс для обработки файлов"""
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Извлекает текст из файла"""
        pass


class PDFHandler(FileHandler):
    """Обработчик PDF файлов"""
    
    def extract_text(self, file_path: str) -> str:
        try:
            import PyPDF2
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except ImportError:
            logger.error("PyPDF2 не установлен")
            raise
        except Exception as e:
            logger.error(f"Ошибка при обработке PDF файла {file_path}: {e}")
            raise


class ExcelHandler(FileHandler):
    """Обработчик Excel файлов"""
    
    def extract_text(self, file_path: str) -> str:
        try:
            import openpyxl
            text = ""
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        text += row_text + "\n"
                text += "\n"
            
            return text.strip()
        except ImportError:
            logger.error("openpyxl не установлен")
            raise
        except Exception as e:
            logger.error(f"Ошибка при обработке Excel файла {file_path}: {e}")
            raise


class WordHandler(FileHandler):
    """Обработчик Word файлов"""
    
    def extract_text(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
            
            return text.strip()
        except ImportError:
            logger.error("python-docx не установлен")
            raise
        except Exception as e:
            logger.error(f"Ошибка при обработке Word файла {file_path}: {e}")
            raise


class FileProcessor:
    """Основной класс для обработки файлов"""
    
    def __init__(self):
        self.handlers = {
            'pdf': PDFHandler(),
            'xlsx': ExcelHandler(),
            'xls': ExcelHandler(),
            'docx': WordHandler(),
            'doc': WordHandler()
        }
    
    def get_handler(self, file_type: str) -> FileHandler:
        """Получает обработчик для конкретного типа файла"""
        handler = self.handlers.get(file_type.lower())
        if not handler:
            raise ValueError(f"Неподдерживаемый тип файла: {file_type}")
        return handler
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Извлекает текст из файла используя соответствующий обработчик"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        handler = self.get_handler(file_type)
        return handler.extract_text(file_path)
    
    def get_supported_types(self) -> List[str]:
        """Возвращает список поддерживаемых типов файлов"""
        return list(self.handlers.keys())
